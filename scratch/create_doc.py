import os
import sys
import subprocess

# Step 1: Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing it...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_document():
    doc = Document()
    
    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style: Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Honey Bee Digital (HBD)\nDashboard Automation System")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # Deep Blue

    # Style: Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("A Non-Technical Guide and Project Report")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    # Add space
    doc.add_paragraph()

    # Section 1: Executive Summary
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Executive Summary")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(31, 78, 121)

    p1 = doc.add_paragraph(
        "Honey Bee Digital manages massive amounts of data key to modern business search and commerce. "
        "This data falls into two primary categories:"
    )
    p1.style.font.name = 'Arial'
    
    p2 = doc.add_paragraph(style='List Bullet')
    r = p2.add_run("Business Listings: ")
    r.bold = True
    p2.add_run("Local directory information across India (names, phones, emails, addresses, cities, states) collected from multiple search portals.")
    
    p3 = doc.add_paragraph(style='List Bullet')
    r = p3.add_run("E-Commerce Products: ")
    r.bold = True
    p3.add_run("Product details, pricing, and category mappings across online retail and quick-commerce marketplaces.")

    doc.add_paragraph(
        "Previously, compiling, cleaning, and validating this data required manual downloads, file management, "
        "and complex Excel formulas. The HBD Dashboard Automation System acts as your central hub to automate these processes. "
        "It connects directly to Google Drive, validates and cleans data in a single click, operates web scrapers in the background, "
        "and generates business reports for team members and clients."
    )

    # Section 2: Core Features Explained
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Core Features Explained")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(18)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(31, 78, 121)

    # Sub-feature A
    sh_a = doc.add_paragraph()
    sh_a_run = sh_a.add_run("A. Automatic Google Drive Sync (The Collector)")
    sh_a_run.font.name = 'Arial'
    sh_a_run.font.size = Pt(14)
    sh_a_run.font.bold = True
    sh_a_run.font.color.rgb = RGBColor(56, 128, 180)

    doc.add_paragraph(
        "Normally, uploading database files requires opening a portal, selecting files, and manually waiting for them to upload. "
        "The system replaces this by checking HBD's designated Google Drive folders every 60 seconds. "
        "When you or your clients upload spreadsheets to Google Drive, the dashboard automatically detects, downloads, parses, "
        "and loads them. You can monitor the ingestion progress directly on the dashboard registry."
    )

    # Sub-feature B
    sh_b = doc.add_paragraph()
    sh_b_run = sh_b.add_run("B. Smart Data Cleaning & Quality Control (The Filter)")
    sh_b_run.font.name = 'Arial'
    sh_b_run.font.size = Pt(14)
    sh_b_run.font.bold = True
    sh_b_run.font.color.rgb = RGBColor(56, 128, 180)

    doc.add_paragraph(
        "Messy directories containing duplicate listings, typos in city names, or invalid contact formats are cleaned automatically:"
    )
    
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Duplicate Removal: ")
    r.bold = True
    p.add_run("The system identifies duplicate entries by checking the unique signature of (Business Name + Phone + Address). Duplicates are safely archived in a review table.")
    
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Location Standardization: ")
    r.bold = True
    p.add_run("Cities, states, and areas are cross-checked against Location_Master_India. Typos are corrected, and listings with unrecognized locations are flagged for manual review.")
    
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Safety Backups (Rollbacks): ")
    r.bold = True
    p.add_run("Before any cleaning operation is applied, the database is backed up. If a clean-up is done incorrectly, you can revert it with one click using the Rollback Registry.")

    # Sub-feature C
    sh_c = doc.add_paragraph()
    sh_c_run = sh_c.add_run("C. Web Scraper Manager (The Explorer)")
    sh_c_run.font.name = 'Arial'
    sh_c_run.font.size = Pt(14)
    sh_c_run.font.bold = True
    sh_c_run.font.color.rgb = RGBColor(56, 128, 180)

    doc.add_paragraph(
        "To gather fresh pricing, categories, and reviews from public portals, the dashboard features a built-in Scraper Manager. "
        "Users can start search-based or link-based scraping tasks for platforms like DMart, Amazon, Zepto, and IndiaMart directly. "
        "Scrapers run safely in the background, updating the system automatically when done."
    )

    # Sub-feature D
    sh_d = doc.add_paragraph()
    sh_d_run = sh_d.add_run("D. Reports & Search Keyword (The Analyst)")
    sh_d_run.font.name = 'Arial'
    sh_d_run.font.size = Pt(14)
    sh_d_run.font.bold = True
    sh_d_run.font.color.rgb = RGBColor(56, 128, 180)

    doc.add_paragraph(
        "The system offers powerful searching and visualization tools. The 'Search Keyword' page lets you query millions of listings "
        "by keyword, city, or category instantly. The reporting panels show source-wise data stats (such as AskLaila, HeyPlaces, and Zomato) "
        "representing completeness, counts, and error summaries."
    )

    # Section 3: How the Data Flows
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. How the Data Flows")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(18)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph(
        "The workflow of data through the system can be divided into four simple steps:"
    )
    
    p = doc.add_paragraph(style='List Number')
    r = p.add_run("Syncing: ")
    r.bold = True
    p.add_run("Messy spreadsheets are picked up from Google Drive folders or gathered by background scrapers.")
    
    p = doc.add_paragraph(style='List Number')
    r = p.add_run("Staging: ")
    r.bold = True
    p.add_run("The raw information is stored in a temporary database table and checked for missing values or duplicates.")
    
    p = doc.add_paragraph(style='List Number')
    r = p.add_run("Cleaning & Review: ")
    r.bold = True
    p.add_run("Users trigger the cleaning pipeline. Standardized entries are formatted, and bad records are moved to review catalogs.")
    
    p = doc.add_paragraph(style='List Number')
    r = p.add_run("Publishing: ")
    r.bold = True
    p.add_run("Cleaned entries are pushed to the main database tables (Listing Master & Product Master), immediately ready for search and reporting.")

    # Section 4: Summary of Tables & Sources
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Summary of Tables & Sources")
    h4_run.font.name = 'Arial'
    h4_run.font.size = Pt(18)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph("The dashboard manages two main types of data collections:")

    # Create Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Shading Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Listing Data Sources (Business Directories)"
    hdr_cells[1].text = "Product Data Sources (E-Commerce Catalog)"
    
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    row_1 = table.rows[1].cells
    row_1[0].text = "Google Maps, JustDial, MagicPin, HeyPlaces, AskLaila, Post Offices, ATMs, Banks, Yellow Pages, NearBuy."
    row_1[1].text = "Amazon, Flipkart (General marketplace catalogs)."

    row_2 = table.rows[2].cells
    row_2[0].text = "College Dunia, Shiksha, SchoolGIS (Educational directories)."
    row_2[1].text = "BigBasket, Blinkit, Zepto, DMart, JioMart, Zomato (Quick commerce & supermarkets)."

    doc.add_paragraph() # Spacer

    # Section 5: Summary of Roles
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Summary of Roles: Who Uses What?")
    h5_run.font.name = 'Arial'
    h5_run.font.size = Pt(18)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(31, 78, 121)

    # Create Roles Table
    table_roles = doc.add_table(rows=5, cols=3)
    table_roles.style = 'Light Shading Accent 1'
    
    hdr_cells_r = table_roles.rows[0].cells
    hdr_cells_r[0].text = "Team Role"
    hdr_cells_r[1].text = "Key Dashboard Section"
    hdr_cells_r[2].text = "What they do"
    
    for cell in hdr_cells_r:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    r1 = table_roles.rows[1].cells
    r1[0].text = "Content Uploaders / Clients"
    r1[1].text = "Google Drive Folders"
    r1[2].text = "Upload directories or sheets to Drive. The automation system imports them automatically."

    r2 = table_roles.rows[2].cells
    r2[0].text = "Data Quality Team"
    r2[1].text = "Cleaning Dashboard & Review Views"
    r2[2].text = "Run 'Dry Runs' to inspect errors, fix unmatched cities/states, and apply cleaning scripts."

    r3 = table_roles.rows[3].cells
    r3[0].text = "Operations Team"
    r3[1].text = "Scrapper Manager"
    r3[2].text = "Launch new search scraping tasks for e-commerce or local map portals."

    r4 = table_roles.rows[4].cells
    r4[0].text = "Project Managers / Clients"
    r4[1].text = "Reports & Search Pages"
    r4[2].text = "Analyze coverage metrics, run keyword queries, and export reports for clients."

    doc.add_paragraph() # Spacer

    # Save
    output_filename = "HBD_Dashboard_Automation_Report.docx"
    doc.save(output_filename)
    print(f"Document saved successfully as: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_document()
